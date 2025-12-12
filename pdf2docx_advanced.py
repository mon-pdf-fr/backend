"""
Advanced PDF to DOCX Converter
High-quality conversion with proper table detection, colors, and layout
"""

import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import List, Dict, Tuple, Optional
from dataclasses import dataclass, field
from collections import defaultdict
import io


@dataclass
class TextSpan:
    """A span of text with consistent formatting"""
    text: str
    x: float
    y: float
    width: float
    height: float
    font_name: str
    font_size: float
    is_bold: bool
    is_italic: bool
    color: Tuple[int, int, int]


@dataclass
class TextLine:
    """A line of text made up of spans"""
    spans: List[TextSpan] = field(default_factory=list)
    y: float = 0.0

    @property
    def text(self) -> str:
        return ''.join(span.text for span in self.spans)

    @property
    def x(self) -> float:
        return self.spans[0].x if self.spans else 0

    @property
    def height(self) -> float:
        return max((s.height for s in self.spans), default=0)


@dataclass
class TableStructure:
    """Detected table structure"""
    headers: List[str]
    rows: List[List[str]]
    y_start: float
    y_end: float
    has_header_bg: bool = False


class AdvancedPDFExtractor:
    """Extract all PDF content with high fidelity"""

    def __init__(self, pdf_path: str):
        self.doc = fitz.open(pdf_path)
        self.pdf_path = pdf_path

    def extract_page_content(self, page_num: int) -> Dict:
        """Extract all content from a page"""
        page = self.doc[page_num]

        content = {
            'text_lines': [],
            'drawings': [],
            'images': [],
            'page_width': page.rect.width,
            'page_height': page.rect.height
        }

        # Extract text with detailed formatting
        text_dict = page.get_text("dict")

        for block in text_dict.get("blocks", []):
            if block.get("type") == 0:  # Text block
                for line in block.get("lines", []):
                    text_line = TextLine()

                    for span in line.get("spans", []):
                        text = span.get("text", "")
                        if not text.strip():
                            continue

                        bbox = span.get("bbox", [0, 0, 0, 0])
                        font = span.get("font", "Arial")
                        size = span.get("size", 12)
                        flags = span.get("flags", 0)
                        color = span.get("color", 0)

                        # Convert color
                        if isinstance(color, int):
                            r = (color >> 16) & 0xFF
                            g = (color >> 8) & 0xFF
                            b = color & 0xFF
                        else:
                            r = g = b = 0

                        text_span = TextSpan(
                            text=text,
                            x=bbox[0],
                            y=bbox[1],
                            width=bbox[2] - bbox[0],
                            height=bbox[3] - bbox[1],
                            font_name=font,
                            font_size=size,
                            is_bold=bool(flags & 16),
                            is_italic=bool(flags & 2),
                            color=(r, g, b)
                        )

                        text_line.spans.append(text_span)

                    if text_line.spans:
                        text_line.y = text_line.spans[0].y
                        content['text_lines'].append(text_line)

        # Extract drawings (for backgrounds and borders)
        drawings = page.get_drawings()
        for drawing in drawings:
            if drawing.get('type') == 'f' and drawing.get('fill'):  # Filled shape
                content['drawings'].append({
                    'rect': drawing.get('rect'),
                    'fill_color': drawing.get('fill'),
                    'type': 'fill'
                })

        # Extract images
        images = page.get_images()
        for img in images:
            try:
                bbox = page.get_image_bbox(img)
                if bbox:
                    xref = img[0]
                    img_data = self.doc.extract_image(xref)
                    content['images'].append({
                        'data': img_data['image'],
                        'bbox': bbox,
                        'ext': img_data.get('ext', 'png')
                    })
            except:
                pass

        return content

    def close(self):
        self.doc.close()


class TableDetector:
    """Detect tables in PDF content"""

    def detect_tables(self, text_lines: List[TextLine]) -> List[TableStructure]:
        """Detect table structures in text lines"""
        tables = []

        # Group lines by Y coordinate with finer precision
        y_groups = defaultdict(list)
        for line in text_lines:
            y_key = round(line.y, 1)  # Group by rounded Y (1 decimal place)
            y_groups[y_key].append(line)

        # Look for table patterns
        sorted_ys = sorted(y_groups.keys())

        i = 0
        while i < len(sorted_ys):
            y = sorted_ys[i]
            current_lines = y_groups[y]

            # Check if this looks like a table header row
            # (multiple columns with bold text)
            if self._is_table_header(current_lines):
                table = self._extract_table(sorted_ys[i:], y_groups)
                if table:
                    tables.append(table)
                    # Skip the rows we've processed
                    i += len(table.rows) + 1
                    continue

            i += 1

        return tables

    def _is_table_header(self, lines: List[TextLine]) -> bool:
        """Check if lines look like a table header"""
        # Can be either:
        # 1. Single line with multiple bold spans (multiple columns)
        # 2. Multiple lines at same Y (PyMuPDF sometimes splits columns into separate "lines")

        if not lines:
            return False

        # Check if all lines are bold
        all_spans = []
        for line in lines:
            all_spans.extend(line.spans)

        if not all_spans:
            return False

        all_bold = all(span.is_bold for span in all_spans)

        # Must have at least 3 columns (more strict - typical for tables)
        if len(all_spans) < 3:
            return False

        # Check for horizontal spacing (columns should be separated)
        if len(all_spans) >= 2:
            # Sort by X position
            sorted_spans = sorted(all_spans, key=lambda s: s.x)

            gaps = []
            for i in range(len(sorted_spans) - 1):
                gap = sorted_spans[i+1].x - (sorted_spans[i].x + sorted_spans[i].width)
                gaps.append(gap)

            # Significant gaps suggest columns
            has_gaps = any(gap > 5 for gap in gaps)

            return all_bold and has_gaps

        return all_bold

    def _extract_table(self, y_keys: List[float], y_groups: Dict) -> Optional[TableStructure]:
        """Extract table structure starting from header"""
        if not y_keys:
            return None

        # First row is header - merge all lines at same Y
        header_y = y_keys[0]
        header_lines = y_groups[header_y]

        # Collect all spans from header lines and sort by X position
        all_header_spans = []
        for line in header_lines:
            all_header_spans.extend(line.spans)

        all_header_spans.sort(key=lambda s: s.x)

        headers = [span.text.strip() for span in all_header_spans]
        num_cols = len(headers)

        # Get column X positions from header
        col_positions = [span.x for span in all_header_spans]

        rows = []
        table_y_start = header_y
        table_y_end = header_y

        # Group nearby Y values as single rows (tolerance of 2 pixels)
        row_groups = []
        current_row_y = None
        current_row_lines = []

        for y in y_keys[1:]:
            lines = y_groups[y]
            if not lines:
                continue

            # Check if this Y is close to current row Y
            if current_row_y is None or abs(y - current_row_y) <= 2.0:
                if current_row_y is None:
                    current_row_y = y
                current_row_lines.extend(lines)
            else:
                # Process previous row
                if current_row_lines:
                    row_groups.append((current_row_y, current_row_lines))
                current_row_y = y
                current_row_lines = list(lines)

            # Stop if we've gone too far
            if y - header_y > 200:
                break

        # Don't forget last row
        if current_row_lines:
            row_groups.append((current_row_y, current_row_lines))

        # Extract rows and merge multi-line cells
        for row_y, row_lines in row_groups:
            # Merge all lines in the row
            merged_line = TextLine()
            for line in row_lines:
                merged_line.spans.extend(line.spans)
            merged_line.y = row_y

            # Check if this is part of the table
            if not self._is_table_row(merged_line, col_positions, num_cols):
                break

            # Extract cell values
            row = self._extract_row_cells(merged_line, col_positions, num_cols)

            # Check if this is a continuation of the previous row
            # (has data only in one column, particularly the description column)
            is_continuation = False
            if rows and len(row) == num_cols:
                non_empty_count = sum(1 for cell in row if cell.strip())
                if non_empty_count == 1:
                    # This might be a multi-line cell continuation
                    # Merge with previous row
                    for i in range(num_cols):
                        if row[i].strip():
                            if rows[-1][i]:
                                rows[-1][i] += ' ' + row[i]
                            else:
                                rows[-1][i] = row[i]
                    is_continuation = True

            if not is_continuation:
                rows.append(row)

            table_y_end = row_y

        if rows:
            return TableStructure(
                headers=headers,
                rows=rows,
                y_start=table_y_start,
                y_end=table_y_end
            )

        return None

    def _is_table_row(self, line: TextLine, col_positions: List[float], num_cols: int) -> bool:
        """Check if line is a table row"""
        # Simple heuristic: line should have text near column positions
        if not line.spans:
            return False

        # Check if any span aligns with column positions (within tolerance)
        tolerance = 30
        for span in line.spans:
            for col_x in col_positions:
                if abs(span.x - col_x) < tolerance:
                    return True

        return False

    def _extract_row_cells(self, line: TextLine, col_positions: List[float], num_cols: int) -> List[str]:
        """Extract cell values from a row"""
        cells = [''] * num_cols

        # Assign each span to the nearest column
        for span in line.spans:
            # Find nearest column
            min_dist = float('inf')
            col_idx = 0
            for i, col_x in enumerate(col_positions):
                dist = abs(span.x - col_x)
                if dist < min_dist:
                    min_dist = dist
                    col_idx = i

            # Add to cell (with space if needed)
            if cells[col_idx]:
                cells[col_idx] += ' '
            cells[col_idx] += span.text.strip()

        return cells


class AdvancedDOCXGenerator:
    """Generate high-quality DOCX with tables and formatting"""

    def __init__(self):
        self.doc = Document()
        self._set_narrow_margins()

    def _set_narrow_margins(self):
        """Set narrow margins for better layout"""
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Cm(1.27)
            section.bottom_margin = Cm(1.27)
            section.left_margin = Cm(1.27)
            section.right_margin = Cm(1.27)

    def add_text_line(self, line: TextLine):
        """Add a text line with proper formatting"""
        para = self.doc.add_paragraph()

        for span in line.spans:
            run = para.add_run(span.text)
            run.font.size = Pt(span.font_size)
            run.font.name = span.font_name
            run.bold = span.is_bold
            run.italic = span.is_italic

            if span.color != (0, 0, 0):
                run.font.color.rgb = RGBColor(*span.color)

    def add_table(self, table_structure: TableStructure):
        """Add a table with proper formatting"""
        num_rows = len(table_structure.rows) + 1  # +1 for header
        num_cols = len(table_structure.headers)

        table = self.doc.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Light Grid Accent 1'

        # Header row
        header_cells = table.rows[0].cells
        for i, header_text in enumerate(table_structure.headers):
            header_cells[i].text = header_text
            # Make header bold
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        # Data rows
        for row_idx, row_data in enumerate(table_structure.rows, start=1):
            cells = table.rows[row_idx].cells
            for col_idx, cell_text in enumerate(row_data):
                if col_idx < len(cells):
                    cells[col_idx].text = cell_text

    def save(self, output_path: str):
        self.doc.save(output_path)


class AdvancedPDF2DOCXConverter:
    """Main converter with advanced features"""

    def __init__(self, pdf_path: str):
        self.pdf_path = pdf_path
        self.extractor = AdvancedPDFExtractor(pdf_path)
        self.table_detector = TableDetector()
        self.generator = AdvancedDOCXGenerator()

    def convert(self, output_path: str, verbose: bool = True):
        """Convert PDF to DOCX with high quality"""
        page_count = self.extractor.doc.page_count

        if verbose:
            print(f"\n{'='*60}")
            print(f"Advanced PDF to DOCX Converter")
            print(f"{'='*60}")
            print(f"Converting {page_count} page(s)...\n")

        for page_num in range(page_count):
            if verbose:
                print(f"Page {page_num + 1}/{page_count}...")

            content = self.extractor.extract_page_content(page_num)
            text_lines = content['text_lines']

            # Detect tables
            tables = self.table_detector.detect_tables(text_lines)

            if verbose and tables:
                print(f"  Found {len(tables)} table(s)")

            # Track which lines are part of tables
            table_y_ranges = [(t.y_start, t.y_end) for t in tables]
            tables_added = set()

            # Process content
            for line in text_lines:
                # Check if this line is part of a table
                in_table_idx = None
                for idx, (y_start, y_end) in enumerate(table_y_ranges):
                    if y_start <= line.y <= y_end:
                        in_table_idx = idx
                        break

                if in_table_idx is not None:
                    # Add table if we haven't yet and this is near the start
                    if in_table_idx not in tables_added:
                        if abs(line.y - tables[in_table_idx].y_start) <= 1.0:
                            self.generator.add_table(tables[in_table_idx])
                            tables_added.add(in_table_idx)
                    # Skip the line (it's part of the table)
                else:
                    # Regular text line
                    self.generator.add_text_line(line)

            # Page break
            if page_num < page_count - 1:
                self.generator.doc.add_page_break()

        # Save
        self.generator.save(output_path)

        if verbose:
            print(f"\n✓ Conversion complete!")
            print(f"  Output: {output_path}")
            print(f"{'='*60}\n")

        self.extractor.close()


def convert_pdf_to_docx(pdf_path: str, output_path: str = None, verbose: bool = True) -> str:
    """
    Convert PDF to DOCX with advanced features

    Features:
    - Automatic table detection
    - Preserves text formatting (bold, italic, colors)
    - Maintains font sizes
    - Proper spacing and layout
    """
    if output_path is None:
        output_path = pdf_path.replace('.pdf', '.docx')

    converter = AdvancedPDF2DOCXConverter(pdf_path)
    converter.convert(output_path, verbose=verbose)

    return output_path
