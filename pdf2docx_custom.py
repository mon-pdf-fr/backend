"""
Custom PDF to DOCX Converter
Built from scratch for high-quality conversion with layout preservation
"""

import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import List, Dict, Tuple, Optional
from dataclasses import dataclass
from collections import defaultdict
import io
from PIL import Image


@dataclass
class TextBlock:
    """Represents a block of text with formatting and position"""
    text: str
    x: float
    y: float
    width: float
    height: float
    font_name: str
    font_size: float
    font_flags: int  # bold, italic, etc.
    color: Tuple[int, int, int]
    page_num: int

    @property
    def is_bold(self) -> bool:
        return bool(self.font_flags & 2**4)  # Bold flag

    @property
    def is_italic(self) -> bool:
        return bool(self.font_flags & 2**1)  # Italic flag

    @property
    def bottom(self) -> float:
        return self.y + self.height

    @property
    def right(self) -> float:
        return self.x + self.width


@dataclass
class ImageBlock:
    """Represents an image with position"""
    image_data: bytes
    x: float
    y: float
    width: float
    height: float
    page_num: int


@dataclass
class TableCell:
    """Represents a table cell"""
    text: str
    row: int
    col: int
    font_size: float
    is_bold: bool


class LayoutAnalyzer:
    """Analyzes PDF layout and groups elements into logical structures"""

    def __init__(self, tolerance_y: float = 3.0, tolerance_x: float = 5.0):
        self.tolerance_y = tolerance_y  # Vertical tolerance for line detection
        self.tolerance_x = tolerance_x  # Horizontal tolerance for alignment

    def group_into_lines(self, blocks: List[TextBlock]) -> List[List[TextBlock]]:
        """Group text blocks into lines based on Y position"""
        if not blocks:
            return []

        # Sort by Y position, then X position
        sorted_blocks = sorted(blocks, key=lambda b: (b.y, b.x))

        lines = []
        current_line = [sorted_blocks[0]]

        for block in sorted_blocks[1:]:
            # Check if block is on the same line (similar Y coordinate)
            if abs(block.y - current_line[0].y) <= self.tolerance_y:
                current_line.append(block)
            else:
                # Start new line
                lines.append(current_line)
                current_line = [block]

        if current_line:
            lines.append(current_line)

        # Sort blocks within each line by X position
        for line in lines:
            line.sort(key=lambda b: b.x)

        return lines

    def merge_line_blocks(self, line: List[TextBlock]) -> TextBlock:
        """Merge multiple text blocks in a line into one"""
        if not line:
            return None

        if len(line) == 1:
            return line[0]

        # Merge text with spaces
        texts = []
        prev_block = None
        for block in line:
            if prev_block:
                # Calculate gap between blocks
                gap = block.x - prev_block.right
                # Add space if there's a gap
                if gap > 2:
                    texts.append(' ')
            texts.append(block.text)
            prev_block = block

        merged_text = ''.join(texts)

        # Use properties from first block
        first = line[0]
        last = line[-1]

        return TextBlock(
            text=merged_text,
            x=first.x,
            y=first.y,
            width=last.right - first.x,
            height=max(b.height for b in line),
            font_name=first.font_name,
            font_size=first.font_size,
            font_flags=first.font_flags,
            color=first.color,
            page_num=first.page_num
        )

    def detect_table(self, lines: List[TextBlock]) -> Optional[List[List[TextBlock]]]:
        """Detect if lines form a table structure"""
        if len(lines) < 2:
            return None

        # Check if lines have similar vertical spacing and column alignment
        # This is a simple heuristic - can be improved

        # For now, return None - tables will be handled as regular text
        # TODO: Implement table detection algorithm
        return None

    def classify_heading(self, block: TextBlock, avg_font_size: float) -> int:
        """Classify text block as heading level (0 = normal text, 1-6 = heading)"""
        # Simple heuristic: larger font = heading
        if block.font_size >= avg_font_size * 1.5:
            return 1  # Heading 1
        elif block.font_size >= avg_font_size * 1.3:
            return 2  # Heading 2
        elif block.font_size >= avg_font_size * 1.15:
            return 3  # Heading 3
        return 0  # Normal text


class PDFExtractor:
    """Extracts content from PDF using PyMuPDF"""

    def __init__(self, pdf_path: str):
        self.doc = fitz.open(pdf_path)
        self.pdf_path = pdf_path

    def extract_text_blocks(self, page_num: int) -> List[TextBlock]:
        """Extract text blocks from a page with formatting info"""
        page = self.doc[page_num]
        blocks = []

        # Get text with formatting details
        text_dict = page.get_text("dict")

        for block in text_dict.get("blocks", []):
            if block.get("type") == 0:  # Text block
                for line in block.get("lines", []):
                    for span in line.get("spans", []):
                        text = span.get("text", "").strip()
                        if not text:
                            continue

                        bbox = span.get("bbox", [0, 0, 0, 0])
                        font = span.get("font", "Arial")
                        size = span.get("size", 12)
                        flags = span.get("flags", 0)
                        color = span.get("color", 0)

                        # Convert color from integer to RGB
                        if isinstance(color, int):
                            r = (color >> 16) & 0xFF
                            g = (color >> 8) & 0xFF
                            b = color & 0xFF
                        else:
                            r = g = b = 0

                        blocks.append(TextBlock(
                            text=text,
                            x=bbox[0],
                            y=bbox[1],
                            width=bbox[2] - bbox[0],
                            height=bbox[3] - bbox[1],
                            font_name=font,
                            font_size=size,
                            font_flags=flags,
                            color=(r, g, b),
                            page_num=page_num
                        ))

        return blocks

    def extract_images(self, page_num: int) -> List[ImageBlock]:
        """Extract images from a page"""
        page = self.doc[page_num]
        images = []

        image_list = page.get_images()

        for img_index, img in enumerate(image_list):
            xref = img[0]
            bbox = page.get_image_bbox(img)

            if bbox and bbox.is_valid:
                try:
                    base_image = self.doc.extract_image(xref)
                    image_data = base_image["image"]

                    images.append(ImageBlock(
                        image_data=image_data,
                        x=bbox.x0,
                        y=bbox.y0,
                        width=bbox.width,
                        height=bbox.height,
                        page_num=page_num
                    ))
                except Exception as e:
                    print(f"Warning: Could not extract image {img_index}: {e}")

        return images

    def get_page_count(self) -> int:
        """Get total number of pages"""
        return len(self.doc)

    def close(self):
        """Close the PDF document"""
        self.doc.close()


class DOCXGenerator:
    """Generates DOCX from extracted content"""

    def __init__(self):
        self.doc = Document()
        self._setup_styles()

    def _setup_styles(self):
        """Setup custom styles for better formatting"""
        styles = self.doc.styles

        # You can customize styles here
        # For now, we'll use default styles

    def add_text_block(self, block: TextBlock, heading_level: int = 0):
        """Add a text block to the document"""
        if heading_level > 0:
            # Add as heading
            para = self.doc.add_heading(block.text, level=heading_level)
        else:
            # Add as paragraph
            para = self.doc.add_paragraph()
            run = para.add_run(block.text)

            # Apply formatting
            run.font.size = Pt(block.font_size)
            run.font.name = block.font_name
            run.bold = block.is_bold
            run.italic = block.is_italic

            # Apply color if not black
            if block.color != (0, 0, 0):
                run.font.color.rgb = RGBColor(*block.color)

    def add_image(self, image_block: ImageBlock, max_width: float = 6.0):
        """Add an image to the document"""
        try:
            # Load image to get dimensions
            img = Image.open(io.BytesIO(image_block.image_data))
            width, height = img.size

            # Calculate appropriate size (maintain aspect ratio)
            aspect_ratio = height / width
            doc_width = min(max_width, image_block.width / 72)  # Convert to inches

            # Add image
            image_stream = io.BytesIO(image_block.image_data)
            self.doc.add_picture(image_stream, width=Inches(doc_width))

        except Exception as e:
            print(f"Warning: Could not add image: {e}")
            # Add placeholder text instead
            self.doc.add_paragraph(f"[Image: {image_block.width:.0f}x{image_block.height:.0f}]")

    def add_table(self, rows: List[List[str]]):
        """Add a table to the document"""
        if not rows:
            return

        table = self.doc.add_table(rows=len(rows), cols=len(rows[0]))
        table.style = 'Light Grid Accent 1'

        for i, row_data in enumerate(rows):
            for j, cell_text in enumerate(row_data):
                table.rows[i].cells[j].text = str(cell_text)

    def save(self, output_path: str):
        """Save the document"""
        self.doc.save(output_path)


class PDF2DOCXConverter:
    """Main converter class that orchestrates the conversion process"""

    def __init__(self, pdf_path: str):
        self.pdf_path = pdf_path
        self.extractor = PDFExtractor(pdf_path)
        self.analyzer = LayoutAnalyzer()
        self.generator = DOCXGenerator()

    def convert(self, output_path: str, verbose: bool = True):
        """Convert PDF to DOCX"""
        page_count = self.extractor.get_page_count()

        if verbose:
            print(f"Converting {page_count} page(s) from PDF to DOCX...")

        all_text_blocks = []

        # Extract all content
        for page_num in range(page_count):
            if verbose:
                print(f"  Processing page {page_num + 1}/{page_count}...")

            # Extract text blocks
            text_blocks = self.extractor.extract_text_blocks(page_num)
            all_text_blocks.extend(text_blocks)

            # Extract images
            images = self.extractor.extract_images(page_num)

            # Combine text and images, sorted by Y position
            elements = []
            elements.extend([(b.y, 'text', b) for b in text_blocks])
            elements.extend([(img.y, 'image', img) for img in images])
            elements.sort(key=lambda x: x[0])

            # Calculate average font size for heading detection
            avg_font_size = sum(b.font_size for b in text_blocks) / len(text_blocks) if text_blocks else 12

            # Group text blocks into lines
            lines = self.analyzer.group_into_lines(text_blocks)

            # Process each line
            for line in lines:
                merged_block = self.analyzer.merge_line_blocks(line)
                if merged_block:
                    # Classify as heading or normal text
                    heading_level = self.analyzer.classify_heading(merged_block, avg_font_size)
                    self.generator.add_text_block(merged_block, heading_level)

            # Add images
            for img in images:
                self.generator.add_image(img)

            # Add page break if not last page
            if page_num < page_count - 1:
                self.generator.doc.add_page_break()

        # Save the document
        self.generator.save(output_path)

        if verbose:
            print(f"✓ Conversion complete: {output_path}")
            print(f"  Extracted {len(all_text_blocks)} text blocks")

        # Cleanup
        self.extractor.close()


def convert_pdf_to_docx(pdf_path: str, output_path: str = None, verbose: bool = True) -> str:
    """
    Convert a PDF file to DOCX format

    Args:
        pdf_path: Path to input PDF file
        output_path: Path to output DOCX file (optional)
        verbose: Print progress messages

    Returns:
        Path to the generated DOCX file
    """
    if output_path is None:
        output_path = pdf_path.replace('.pdf', '.docx')

    converter = PDF2DOCXConverter(pdf_path)
    converter.convert(output_path, verbose=verbose)

    return output_path
